import urllib.error

import newspaper
from docx import Document
from bs4 import BeautifulSoup
from newspaper import Article
from googlesearch import search

BANNED_WEBSITES = ['www.youtube.com', 'www.facebook.com', 'www.instagram.com', 'www.twitter.com', 'uk.linkedin.com',
                   'www.linkedin.com']


class ArticlesScraper:
    def __init__(self, argument=None, number_of_links=10, progress_bar=None):
        self.argument = argument
        self.number_of_links = number_of_links
        self.progress_bar = progress_bar
        if type(self.argument) is str:
            self.links = self.get_google_links(self.argument)
        else:
            self.links = self.argument
        if not self.links:
            self.result = False
        else:
            self.result = self.concatenate_articles_into_one_docx_file()

    def concatenate_articles_into_one_docx_file(self):
        document = Document()
        document.add_heading("Search for {}".format(self.argument), 0)
        for link in self.links:
            self.progress_bar.emit()
            document = self.add_single_article_text(link, document)
        return document

    def add_single_article_text(self, url, document):
        try:
            url = self.reformat_url(url)
            article_info = Article(url)
            article_info.download()
            article_info.parse()
            soup = BeautifulSoup(article_info.html, 'html.parser')
            p = document.add_paragraph(url + "\n\n")
            p.add_run("TITLE: ").bold = True
            p.add_run(article_info.title + "\n")
            p.add_run("META DESCRIPTION: ").bold = True
            p.add_run(article_info.meta_description + "\n")
            p.add_run("META KEYWORDS: ").bold = True
            for keyword in article_info.meta_keywords:
                if keyword != '':
                    p.add_run(keyword + ',')
            p.add_run('\nH1').bold = True
            p.add_run(self.unpack_list_of_strings(soup, 'h1'))
            p.add_run('\nH2').bold = True
            p.add_run(self.unpack_list_of_strings(soup, 'h2'))
            p.add_run('\nBOLD/STRONG: ').bold = True
            p.add_run(self.unpack_list_of_strings(soup, 'strong'))
            p.add_run(self.unpack_list_of_strings(soup, 'b'))
            p.add_run('\nItalic: : ').bold = True
            p.add_run(self.unpack_list_of_strings(soup, 'em'))
            p.add_run(self.unpack_list_of_strings(soup, 'i'))
            p.add_run("\nTEXT: ").bold = True
            p.add_run(article_info.text + '\n\n\n')
            return document
        except newspaper.ArticleException:
            document.add_paragraph(url + "\n" + "Access to this page has been denied!" + '\n\n\n')
            return document

    def get_google_links(self, query):
        try:
            links = []
            used_websites = []
            for j in search(query, num=50, stop=50, pause=2):
                if 'https' in j:
                    host_name = j.split('https://')[1].split('/')[0]
                else:  # if url starts with http only
                    host_name = j.split('http://')[1].split('/')[0]
                if host_name not in used_websites and host_name not in BANNED_WEBSITES:
                    links.append(j)
                    used_websites.append(host_name)
                if len(links) >= self.number_of_links:
                    break
            return links
        except urllib.error.URLError:
            return False

    @staticmethod
    def reformat_url(url):
        url = url.replace('\n', '')
        if url[0] == '\'' or url[0] == '\"':
            url = url[1:]
        if url[-1] == '\'' or url[-1] == '\"':
            url = url[:-1]
        return url

    @staticmethod
    def unpack_list_of_strings(soup, html_element):
        sub_class_str = ''
        for phrase in soup.find_all(html_element):
            if phrase.get_text() != '':
                sub_class_str += phrase.get_text() + ', '
        return sub_class_str


if __name__ == '__main__':
    h = ['https://en.wikipedia.org/wiki/Niagara_Falls', 'https://simple.wikipedia.org/wiki/Niagara_Falls',
         'https://www.britannica.com/place/Niagara-Falls-waterfall-North-America',
         'https://www.niagarafallsstatepark.com/', 'https://www.niagarafallstourism.com/',
         'https://www.niagarafallslive.com/',
         'https://www.world-of-waterfalls.com/waterfalls/eastern-us-niagara-falls/',
         'https://earth.esa.int/web/earth-watching/image-of-the-week/content/-/article/niagara-falls-2018/index.html',
         'https://pixabay.com/images/search/niagara%20falls/', 'https://www.niagarafallsusa.com/']
    a = ArticlesScraper(h)
    document = Document()
    document.add_paragraph(a.result)
    document.save(r'C:\Users\user\PycharmProjects\googlescraper\links.docx')
